import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from pypdf import PdfReader
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)
from reportlab.lib.styles import getSampleStyleSheet

from rag import (
    ask_resume,
    analyze_resume,
    skill_growth,
    generate_resume,
    jd_match,
    extract_resume_data
)

# ==========================================
# PAGE CONFIG
# ==========================================
st.markdown("""
<style>

[data-testid="stAppViewContainer"]{
    background: linear-gradient(
        135deg,
        #0f172a,
        #111827
    );
}

.main-title{
    font-size:42px;
    font-weight:700;
    color:white;
    text-align:center;
}

h1,h2,h3,h4,h5,h6{
    color:white !important;
}

p,label{
    color:white !important;
}

.card{
    background:#1e293b;
    padding:20px;
    border-radius:20px;
    border:1px solid #334155;
    box-shadow:0 5px 15px rgba(0,0,0,0.3);
}

.section-title{
    color:#38bdf8;
    font-size:22px;
    font-weight:600;
}

.stButton button{
    width:100%;
    background:#2563eb;
    color:white;
    border:none;
    border-radius:12px;
    height:50px;
    font-size:18px;
    font-weight:600;
}

.stButton button:hover{
    background:#1d4ed8;
}

.stTextInput input,
.stTextArea textarea{
    border-radius:12px;
}

[data-testid="metric-container"]{
    background:#1e293b;
    border:1px solid #334155;
    padding:20px;
    border-radius:15px;
}

[data-testid="stDataFrame"]{
    border-radius:15px;
    border:1px solid #334155;
}

</style>
""", unsafe_allow_html=True)

# ==========================================
# HEADER
# ==========================================

st.title("🚀 AI Resume Suite")

col1, col2, col3 = st.columns(3)

with col1:
    st.metric("Resume Builder", "AI")

with col2:
    st.metric("ATS Checker", "100%")

with col3:
    st.metric("Career Growth", "Enabled")

# ==========================================
# DOWNLOAD HELPERS
# ==========================================

def create_docx(resume_text):

    doc = Document()

    doc.add_heading(
        "Resume",
        level=1
    )

    doc.add_paragraph(
        resume_text
    )

    buffer = BytesIO()

    doc.save(buffer)

    buffer.seek(0)

    return buffer


def create_pdf(resume_text):

    buffer = BytesIO()

    pdf = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    content = []

    for line in resume_text.split("\n"):

        if line.strip():

            content.append(
                Paragraph(
                    line,
                    styles["BodyText"]
                )
            )

            content.append(
                Spacer(1, 4)
            )

    pdf.build(content)

    buffer.seek(0)

    return buffer

# ==========================================
# TABS
# ==========================================

tab1, tab2, tab3, tab4, tab5 = st.tabs(
    [
        "💬 Resume Chat",
        "📊 ATS Analysis",
        "📈 Skill Growth",
        "📝 Resume Builder",
        "🎯 JD Match"
    ]
)

# ==========================================
# RESUME CHAT
# ==========================================

with tab1:

    st.subheader(
        "Resume Chatbot"
    )

    question = st.text_input(
        "Ask anything about your resume"
    )

    if st.button(
        "Ask Resume"
    ):

        if question:

            with st.spinner(
                "Searching..."
            ):

                answer = ask_resume(
                    question
                )

            st.success(
                "Answer Generated"
            )

            st.write(answer)

# ==========================================
# ATS ANALYSIS
# ==========================================

with tab2:

    st.subheader(
        "ATS Resume Checker"
    )

    if st.button(
        "Analyze Resume"
    ):

        with st.spinner(
            "Analyzing..."
        ):

            result = analyze_resume()

        st.markdown(result)

# ==========================================
# SKILL GROWTH
# ==========================================

with tab3:

    st.subheader(
        "Career Growth Advisor"
    )

    if st.button(
        "Generate Growth Plan"
    ):

        with st.spinner(
            "Analyzing Skills..."
        ):

            result = skill_growth()

        st.markdown(result)

# ==========================================
# RESUME BUILDER
# ==========================================

with tab4:

    st.subheader(
        "AI Resume Builder"
    )

    # ----------------------------------
    # PROFILE IMAGE
    # ----------------------------------

    profile_pic = st.file_uploader(
        "Upload Profile Picture",
        type=["jpg", "jpeg", "png"]
    )

    if profile_pic:

        st.image(
            profile_pic,
            width=150
        )

    # ----------------------------------
    # RESUME UPLOAD
    # ----------------------------------

    uploaded_resume = st.file_uploader(
        "Upload Existing Resume",
        type=["pdf"]
    )

    auto_name = ""
    auto_email = ""
    auto_phone = ""

    if uploaded_resume:

        reader = PdfReader(
            uploaded_resume
        )

        resume_text = ""

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:
                resume_text += page_text

        extracted = extract_resume_data(
            resume_text
        )

        auto_name = extracted["name"]
        auto_email = extracted["email"]
        auto_phone = extracted["phone"]

        st.success(
            "Resume Parsed Successfully"
        )

    # ----------------------------------
    # TEMPLATE
    # ----------------------------------

    template = st.selectbox(
    "Select Resume Template",
    [
        "Modern",
        "Professional",
        "Harvard"
    ]
)

    # ----------------------------------
    # FORM
    # ----------------------------------

    col1, col2 = st.columns(2)

    with col1:

        name = st.text_input(
            "Full Name",
            auto_name
        )

        email = st.text_input(
            "Email",
            auto_email
        )

        phone = st.text_input(
            "Phone",
            auto_phone
        )

        role = st.text_input(
            "Target Role",
            "Computer Vision Engineer"
        )

    with col2:

        education = st.text_area(
            "Education"
        )

        certifications = st.text_area(
            "Certifications"
        )

    skills = st.text_area(
        "Skills"
    )

    experience = st.text_area(
        "Experience"
    )

    projects = st.text_area(
        "Projects"
    )

    # ----------------------------------
    # PREVIEW TABLE
    # ----------------------------------

    st.subheader(
        "Preview"
    )

    preview_df = pd.DataFrame(
        {
            "Field": [
                "Name",
                "Email",
                "Phone",
                "Role",
                "Template"
            ],
            "Value": [
                name,
                email,
                phone,
                role,
                template
            ]
        }
    )

    st.dataframe(
        preview_df,
        use_container_width=True
    )

    # ----------------------------------
    # GENERATE RESUME
    # ----------------------------------

    if st.button(
        "Generate Resume"
    ):

        with st.spinner(
            "Generating Resume..."
        ):
# ----------------------------------
# LIVE RESUME PREVIEW
# ----------------------------------
#    # ----------------------------------
    # LIVE RESUME PREVIEW
    # ----------------------------------

    st.subheader("📄 Live Resume Preview")

    resume_html = f"""
    <div style="
        background:white;
        padding:40px;
        border-radius:20px;
        box-shadow:0 5px 20px rgba(0,0,0,0.2);
        color:black;
    ">
        <div style="text-align:center;">
            <h1>{name}</h1>
            <p>{email} | {phone}</p>
            <h3 style="color:#2563eb;">{role}</h3>
        </div>

        <hr>

        <h3 style="color:#2563eb;">Skills</h3>
        <p>{skills}</p>

        <h3 style="color:#2563eb;">Experience</h3>
        <p>{experience}</p>

        <h3 style="color:#2563eb;">Projects</h3>
        <p>{projects}</p>

        <h3 style="color:#2563eb;">Education</h3>
        <p>{education}</p>

        <h3 style="color:#2563eb;">Certifications</h3>
        <p>{certifications}</p>
    </div>
    """

    st.markdown(
        resume_html,
        unsafe_allow_html=True
    )

    # ----------------------------------
    # GENERATE RESUME
    # ----------------------------------

    if st.button("Generate Resume"):

        with st.spinner("Generating Resume..."):

            resume = generate_resume(
                template,
                name,
                email,
                phone,
                role,
                skills,
                experience,
                projects,
                education,
                certifications
            )

        st.success("Resume Generated Successfully")

        st.markdown(
            resume_html,
            unsafe_allow_html=True
        )

        # DOCX Download
        docx_file = create_docx(resume)

        st.download_button(
            "📄 Download DOCX",
            data=docx_file,
            file_name="resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # PDF Download
        pdf_file = create_pdf(resume)

        st.download_button(
            "📕 Download PDF",
            data=pdf_file,
            file_name="resume.pdf",
            mime="application/pdf"
        )
# ==========================================
# JD MATCH
# ==========================================

with tab5:

    st.subheader("Job Description Match")

    job_description = st.text_area(
        "Paste Job Description"
    )

    if st.button(
        "Match Resume"
    ):

        if job_description:

            with st.spinner(
                "Calculating Match..."
            ):

                result = jd_match(
                    job_description
                )

            st.markdown(
                result
            )